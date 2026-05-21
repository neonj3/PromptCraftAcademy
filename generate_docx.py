from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "PromptCraft-Academy-Blueprint.docx"


COURSE_SOURCES = [
    {
        "name": "Learn Prompting - Introduction to Prompt Engineering",
        "summary": (
            "Beginner-oriented prompt engineering course covering prompt basics, model limitations, "
            "hallucinations, bias, and responsible AI practice."
        ),
        "url": "https://stage.learnprompting.org/courses/introduction-to-prompt-engineering",
    },
    {
        "name": "Microsoft Generative AI for Beginners",
        "summary": (
            "A 21-lesson open course that mixes concept lessons with build lessons and includes code samples "
            "plus extra learning resources."
        ),
        "url": "https://github.com/microsoft/generative-ai-for-beginners",
    },
    {
        "name": "Hugging Face LLM Course",
        "summary": (
            "Free course on LLMs and NLP foundations, using the Hugging Face ecosystem and extending into "
            "advanced topics like fine-tuning and reasoning models."
        ),
        "url": "https://huggingface.co/learn/llm-course/chapter1/1",
    },
    {
        "name": "Google Skills - Beginner: Introduction to Generative AI",
        "summary": (
            "A five-activity beginner learning path managed by Google Cloud that covers GenAI basics, "
            "large language models, and responsible AI principles."
        ),
        "url": "https://www.skills.google/paths/118?locale=en",
    },
]


GAME_MODULES = [
    ("Prompt Foundations", "Role, task, audience, and output format."),
    ("Context and Constraints", "Add background, scope, tone, and limits."),
    ("Output Design and Evaluation", "Request structure and score responses with a rubric."),
    ("Hallucinations and Verification", "Separate facts, uncertainty, and evidence checks."),
    ("Few-shot Prompting and Chaining", "Use examples and split hard tasks into smaller steps."),
    ("Grounding and RAG Basics", "Answer from source material and avoid guessing."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_bullet_list(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    style_name = f"Heading {level}"
    paragraph.style = style_name
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    if level == 1:
      run.font.size = Pt(16)
    elif level == 2:
      run.font.size = Pt(14)
    else:
      run.font.size = Pt(12)
    run.bold = True
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    widths = [Inches(2.2), Inches(4.3)]
    if len(headers) == 3:
        widths = [Inches(1.9), Inches(2.0), Inches(2.6)]

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].width = widths[index]
        set_cell_shading(header_cells[index], "EAF4F1")
        set_cell_text(header_cells[index], header, bold=True)

    for row_data in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_data):
            row_cells[index].width = widths[index]
            set_cell_text(row_cells[index], value)

    for row in table.rows:
        row.height = None
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "bottom", "start", "end"):
                element = tc_mar.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    tc_mar.append(element)
                element.set(qn("w:w"), "100" if side in ("top", "bottom") else "120")
                element.set(qn("w:type"), "dxa")

    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for heading_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = styles[heading_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(6 if heading_name != "Heading 1" else 8)

    if "Subtitle" not in [style.name for style in styles]:
        subtitle = styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)


def build_doc():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PromptCraft Academy")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Consolidated project brief, curriculum mapping, deployment plan, and build notes for a zero-cost GenAI learning game."
    )

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run(f"Date: {date.today().isoformat()}   |   Format: static web app + Word brief")
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(10.5)
    meta_run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. Project overview", level=1)
    doc.add_paragraph(
        "PromptCraft Academy is a clean, beginner-friendly browser game that teaches practical generative AI skills "
        "without requiring paid APIs, a backend, or a complex setup. The project is designed so an amateur can learn "
        "core GenAI concepts and immediately apply them in real workflows such as studying, content creation, resume writing, "
        "research, and internal knowledge support."
    )

    add_heading(doc, "2. Product goals", level=1)
    add_bullet_list(
        doc,
        [
            "Teach GenAI from basic to intermediate level through short lessons and replayable missions.",
            "Keep the interface clean, calm, and readable on desktop and mobile.",
            "Make the game fully static so it can be deployed for free on GitHub Pages, Netlify, or Cloudflare Pages.",
            "Focus on practical prompt writing, response evaluation, verification, grounding, and responsible AI usage.",
            "Use simple mission scoring instead of live model calls so the product remains zero-cost.",
        ],
    )

    add_heading(doc, "3. Learning sources used", level=1)
    source_rows = [(source["name"], source["summary"]) for source in COURSE_SOURCES]
    add_table(doc, ["Source", "How it shaped the game"], source_rows)
    source_note = doc.add_paragraph()
    source_note.paragraph_format.space_before = Pt(4)
    note_run = source_note.add_run(
        "Reference links: "
        + "; ".join(f"{source['name']} - {source['url']}" for source in COURSE_SOURCES)
    )
    note_run.font.name = "Arial"
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "4. Curriculum progression", level=1)
    curriculum_rows = [(str(index + 1), name, summary) for index, (name, summary) in enumerate(GAME_MODULES)]
    add_table(doc, ["Module", "Topic", "Outcome"], curriculum_rows)

    add_heading(doc, "5. Game loop", level=1)
    add_bullet_list(
        doc,
        [
            "Player opens a module and reads the lesson cards: core concept, why it matters, and a practical pattern.",
            "The mission shows a scenario, goal, starter prompt, and a list of success criteria.",
            "Player rewrites the prompt in the workbench.",
            "The scoring engine checks for the presence of role, context, constraints, structure, and safety signals depending on the module.",
            "On passing, the player earns XP, unlocks a reflection takeaway, and stores progress in browser localStorage.",
        ],
    )

    add_heading(doc, "6. Technical architecture", level=1)
    add_table(
        doc,
        ["Layer", "Decision"],
        [
            ("Frontend", "Plain HTML, CSS, and JavaScript for zero dependency deployment."),
            ("Persistence", "Browser localStorage for XP, completed missions, and streak state."),
            ("Scoring", "Heuristic keyword and structure checks instead of live model evaluation."),
            ("Hosting", "Free static hosting only; no backend required."),
        ],
    )

    add_heading(doc, "7. File structure", level=1)
    add_bullet_list(
        doc,
        [
            "index.html - app shell, learning panels, mission workbench, and progress layout.",
            "styles.css - clean responsive UI, visual system, and interaction styles.",
            "app.js - module content, scoring logic, mission switching, and local progress saving.",
            "README.md - local run instructions, deployment steps, and project summary.",
            "PromptCraft-Academy-Blueprint.docx - this consolidated project document.",
        ],
    )

    add_heading(doc, "8. Why the app is beginner-friendly", level=1)
    add_bullet_list(
        doc,
        [
            "Uses plain language instead of deep technical jargon.",
            "Explains each concept before asking the player to practice it.",
            "Connects each mission to a relatable real-life use case.",
            "Gives feedback on what is missing rather than only giving a numeric score.",
            "Keeps the prompt formula visible so learners can reuse it outside the game.",
        ],
    )

    add_heading(doc, "9. Zero-cost deployment instructions", level=1)
    doc.add_paragraph("Recommended path: GitHub Pages, because the project is static and does not need a build step.")
    add_bullet_list(
        doc,
        [
            "Create a GitHub repository and upload index.html, styles.css, and app.js to the root.",
            "Open repository Settings, then Pages.",
            "Choose Deploy from a branch.",
            "Select the main branch and the root folder.",
            "Save and wait for GitHub to publish the site URL.",
        ],
    )
    doc.add_paragraph("Alternative free options: Netlify and Cloudflare Pages. Both can deploy the project with no build command.")

    add_heading(doc, "10. How Codex can help deploy it", level=1)
    add_bullet_list(
        doc,
        [
            "Ask Codex to publish the static site to GitHub Pages without adding a backend.",
            "Ask Codex to verify that the project remains zero-cost and fully static.",
            "Ask Codex to customize the curriculum, visuals, or number of missions before deployment.",
        ],
    )

    add_heading(doc, "11. Suggested future improvements", level=1)
    add_bullet_list(
        doc,
        [
            "Add image-prompting missions and multimodal lessons.",
            "Add a downloadable teacher guide or worksheet pack.",
            "Add richer rubric checks for chain-of-thought alternatives, structured outputs, and prompt safety.",
            "Add a content pack of grounded source documents so learners can practice RAG-style prompting.",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
